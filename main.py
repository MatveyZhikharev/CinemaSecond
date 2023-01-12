from docx import Document
import xlsxwriter


def docx(cinema):
    document = Document()

    document.add_heading(f'Отчёт на кинотеатр на {cinema.adress}', 0)

    table = document.add_table(rows=len(cinema.halls) + 1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер'
    hdr_cells[1].text = 'Количество свободных билетов'
    hdr_cells[2].text = 'Количество проданных билетов'
    print(cinema.halls)

    for i in range(0, len(cinema.halls)):
        hdr_cells = table.rows[i + 1].cells
        hdr_cells[0].text = str(i + 1)
        hdr_cells[1].text = str(cinema.halls[i].free_arms())
        hdr_cells[2].text = str(cinema.halls[i].not_free_arms())

    document.save('info.docx')


def xlsx(cinema):
    workbook = xlsxwriter.Workbook('info.xlsx')
    worksheet = workbook.add_worksheet()

    worksheet.write(0, 0, 'Номер')
    worksheet.write(0, 1, 'Количество свободных билетов')
    worksheet.write(0, 2, 'Количество проданных билетов')

    data = [(str(cinema.halls[i].free_arms()), str(cinema.halls[i].not_free_arms()))
            for i in range(0, len(cinema.halls))]

    for row, (first, second, third) in enumerate(data):
        worksheet.write(row + 1, 0, first)
        worksheet.write(row + 1, 1, second)
        worksheet.write(row + 1, 1, third)

    workbook.close()


class Cinema_network:
    def __init__(self, name, *cinemas):  # на вход имя сети и кинотеатры входящие в нее
        self.name = name
        if cinemas:
            self.cinemas = list(cinemas)
        else:
            self.cinemas = []

    def append(self, item):
        self.cinemas.append(item)

    def __str__(self):
        return "\n".join(cinema.__str__() for cinema in self.cinemas)

    def __repr__(self):
        return f"{self.__class__.__name__}('{self.name}', {self.cinemas})"


class Cinema:
    counter = 0

    def __init__(self, adress, *halls):
        self.adress = adress
        if halls:
            self.halls = list(halls)
        else:
            self.halls = []

    def append(self, item):
        self.halls.append(item)

    def __str__(self):
        return f"Кинотеатр на {self.adress} {', '.join(hall.__str__() for hall in self.halls)}"

    def __repr__(self):
        return f"{self.__class__.__name__}('{self.adress}', {self.halls})"


class Hall:
    def __init__(self, num, *armchairs):
        self.num = num
        if armchairs:
            self.armchairs = list(armchairs)
            self.full = 0
        else:
            self.armchairs = [[]]
            self.full = 1

    def free_arms(self):
        return "".join(map(str, self.armchairs)).count("0")

    def not_free_arms(self):
        return "".join(map(str, self.armchairs)).count("1")

    def edit(self, item):
        self.armchairs = [[int(armch) for armch in line] for line in item.split(";")]
        self.full = "0" in item

    def choose_place(self, *rowcol):
        if not self.full:
            self.armchairs[rowcol[0]][rowcol[1]] = "1"
            print(f"{rowcol[1]} место на {rowcol[0]} ряду забронировано")
            if "0" not in str(self.armchairs):
                self.full = 1
        else:
            print('Все места заняты')

    def __str__(self):
        hall = "\n".join("|".join(str(el) for el in row) for row in self.armchairs)
        return f"Зал, номер {self.num}, {'Все места заняты' if self.full else 'Есть свободный места'}\n{hall}"

    def __repr__(self):
        return f"{self.__class__.__name__}('{self.num}', {self.armchairs})"


def print_comands():
    print(f"Команды:\n{'1'}{'Выбрать кинотеатр':>50}\n{'2'}{'Выбрать зал':>50}\n{'3'}{'Выбрать сидение':>50}"
          f"\n{'4'}{'Стоп':>50}")
    print(f"Продвинутые команды:"
          f"\n{'5'}{'Добавить кинотеатр':>50}"
          f"\n{'6'}{'Добавить зал':>50}"
          f"\n{'7'}{'Изменить сидения':>50}"
          f"\n{'8'}{'Изменить доступность сидения':>50}"
          f"\n{'9'}{'Отчёт Microsoft Word':>50}"
          f"\n{'10'}{'Отчёт Microsoft Excel':>50}")


if __name__ == "__main__":
    Cinema_network(input("Напиши название сети кинотетров"))
    cinemas = []
    last_cinema = ""
    last_hall = ""
    while True:
        try:
            print_comands()
            command = input().lower().strip()
            if command in ("добавить кинотеатр", "5"):
                cinemas.append(last_cinema := Cinema(input("Введите адрес:")))
            elif command in ("добавить зал", "6"):
                if not last_cinema:
                    print("Выберете номер кинотеатра из списка:", *cinemas, sep="\n")
                    last_cinema = cinemas[int(input()) - 1]
                last_cinema.halls.append(Hall(input("Введите название:")))
            elif command in ("изменить сидения", "7"):
                if not last_cinema:
                    print("Выберете номер кинотеатра из списка:", *cinemas, sep="\n")
                    last_cinema = cinemas[int(input()) - 1]
                if not last_hall:
                    print("Выберете зал:", *last_cinema.halls)
                    last_hall = last_cinema.halls[int(input()) - 1]
                print("Напечатайте схему сидений в формате: каждый ряд на новой строке, занятое кресло - 1, "
                      "свободное - 0, ряды разделяйте ;")
                last_hall.edit(input())
            elif command in ("изменить доступность сидения", "8"):
                if not last_cinema:
                    print("Выберете номер кинотеатра из списка:", *cinemas, sep="\n")
                    last_cinema = cinemas[int(input()) - 1]
                if not last_hall:
                    print("Выберете зал:", *last_cinema.halls)
                    last_hall = last_cinema.halls[int(input()) - 1]
                print("Выедите ряд и номер сидения через запятую:")
                print(type(last_hall))
                print(last_hall)
            elif command in ("выбрать кинотеатр", "1"):
                print("Выберете номер кинотеатра из списка:", *cinemas, sep="\n")
                last_cinema = cinemas[int(input()) - 1]
            elif command in ("выбрать зал", "2"):
                if not last_cinema:
                    print("Сперва выберете кинотетр")
                    continue
                print("Выберете номер зала из списка:", last_cinema.halls, sep="\n")
                last_hall = cinemas[int(input()) - 1]
            elif command in ("выбрать сидение", "3"):
                if not last_hall:
                    print("Сперва выберете зал")
                    continue
                print("Выберете номер сидения из списка:", last_hall, sep="\n")
                n, m = map(lambda x: int(x) - 1, input().split())
                if not last_hall.armchairs[n][m]:
                    last_hall.armchairs[n][m] = 1
                    print("место выбрано")
            elif command in ("cтоп", "4"):
                break
            elif command in ("отчёт microsoft word", "9"):
                if not last_cinema:
                    print("Сперва выберите кинотеатр")
                elif last_cinema.halls[0].armchairs:
                    docx(last_cinema)
                else:
                    print("Добавьте сидения")
            elif command in ("отчёт microsoft excel", "10"):
                if not last_cinema:
                    print("Сперва выберите кинотеатр")
                elif last_cinema.halls[0].armchairs:
                    xlsx(last_cinema)
                else:
                    print("Добавьте сидения")
            else:
                print("Я тебя не понял повтори пожалуйста")
        except Exception as e:
            print(e)
